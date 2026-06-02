"""DOCX template engine.

Generates formatted Word documents from structured report data + a template config.
Config defines document structure: sections, headings, tables, paragraphs.
No domain knowledge — purely mechanical formatting.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    raise ImportError("python-docx is required. Install with: pip install python-docx")


def generate(
    report: dict[str, Any],
    template_config: dict,
    output_path: str | Path,
    *,
    metadata: dict[str, Any] | None = None,
) -> Path:
    """Generate a formatted DOCX from report data + template config.

    Template config structure:
        page:
          size: A4 | letter
          margins: {top, bottom, left, right}  # cm
          font: <str>
          font_size: <int>  # pt

        title:
          text: <str>
          size: <int>
          color: <hex>
          alignment: center | left | right

        sections:
          - heading: <str>
            heading_style: <str>  # styled-header | plain
            heading_color: <hex>
            elements:
              - type: paragraph
                text: <str>
                bold: true/false
                format: <str>  # python format string with {value}
                source: <str>  # dot-path to value in report data

              - type: key_value
                source: <str>
                label_style: bold | normal

              - type: table
                source: <str>   # dot-path to list[dict] in report data
                columns:
                  - field: <str>
                    header: <str>
                    width_pct: <int>

              - type: list
                source: <str>   # dot-path to list in report data
                item_format: <str>

              - type: section_departments
                source: <str>   # dot-path to dict {dept: {stats: ..., comments: [...]}}
                counters_label: <str>  # "Counters" etc
                comments_label: <str>  # "Comments" etc

        footer:
          text: <str>
          include_timestamp: true/false
    """
    cfg = template_config
    page_cfg = cfg.get("page", {})
    meta = metadata or {}

    doc = Document()

    # Page setup
    sec = doc.sections[0]
    if page_cfg.get("size") == "A4":
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
    margins = page_cfg.get("margins", {})
    sec.top_margin = Cm(margins.get("top", 2))
    sec.bottom_margin = Cm(margins.get("bottom", 2))
    sec.left_margin = Cm(margins.get("left", 2.5))
    sec.right_margin = Cm(margins.get("right", 2))

    # Default font
    font_name = page_cfg.get("font", "Calibri")
    font_size = page_cfg.get("font_size", 10)
    doc.styles["Normal"].font.name = font_name
    doc.styles["Normal"].font.size = Pt(font_size)

    # Title
    title_cfg = cfg.get("title", {})
    t = doc.add_paragraph()
    t.alignment = _alignment(title_cfg.get("alignment", "center"))
    title_text = title_cfg.get("text", "Report")
    # Inject metadata
    for mk, mv in meta.items():
        title_text = title_text.replace(f"{{{mk}}}", str(mv))
    r = t.add_run(title_text)
    r.font.size = Pt(title_cfg.get("size", 18))
    r.bold = True
    r.font.color.rgb = _color(title_cfg.get("color", "#003366"))

    # Subtitle
    if subtitle := cfg.get("subtitle"):
        s = doc.add_paragraph()
        s.alignment = _alignment(subtitle.get("alignment", "center"))
        r = s.add_run(subtitle["text"])
        r.font.size = Pt(subtitle.get("size", 13))
        r.bold = True
        r.font.color.rgb = _color(subtitle.get("color", "#003366"))
        s.paragraph_format.space_after = Pt(10)

    # Metadata section
    if meta_fields := cfg.get("meta_fields"):
        for mf in meta_fields:
            _add_para(doc, mf.get("label", ""), bold=True, size=font_size)
            val = _resolve(report, mf.get("source", ""))
            if isinstance(val, list):
                for item in val:
                    _add_para(doc, str(item), size=font_size - 1, after=2)
            else:
                _add_para(doc, str(val or mf.get("default", "")), size=font_size, after=6)

    # Sections
    for section in cfg.get("sections", []):
        _render_section(doc, section, report, font_size)

    # Footer
    footer_cfg = cfg.get("footer", {})
    doc.add_paragraph()
    if footer_text := footer_cfg.get("text"):
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(footer_text)
        fr.font.size = Pt(9)
        fr.italic = True
        fr.font.color.rgb = RGBColor(128, 128, 128)

    if footer_cfg.get("include_timestamp", True):
        fp2 = doc.add_paragraph()
        fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generator = meta.get("generator", "excel2docx")
        fr2 = fp2.add_run(f"Auto-generated: {datetime.now():%Y-%m-%d %H:%M} | {generator}")
        fr2.font.size = Pt(8)
        fr2.font.color.rgb = RGBColor(160, 160, 160)

    output_path = Path(output_path)
    doc.save(str(output_path))
    return output_path


# ── Internal helpers ──────────────────────────────────────────────


def _render_section(doc, section_cfg: dict, data: dict, base_size: int):
    """Render one section."""
    heading = section_cfg.get("heading", "")
    heading_style = section_cfg.get("heading_style", "styled-header")

    if heading_style == "styled-header":
        _add_styled_header(doc, heading, section_cfg.get("heading_color", "#2F5496"))
    else:
        _add_para(doc, heading, bold=True, size=base_size + 1)

    for element in section_cfg.get("elements", []):
        el_type = element.get("type", "paragraph")
        source = element.get("source", "")

        if el_type == "paragraph":
            text = element.get("text", "")
            if "{" in text:
                val = _resolve(data, source) if source else ""
                text = text.format(value=val)
            _add_para(doc, text,
                     bold=element.get("bold", False),
                     size=element.get("size", base_size))

        elif el_type == "key_value":
            items = _resolve(data, source) if source else {}
            if isinstance(items, dict):
                for k, v in items.items():
                    p = doc.add_paragraph()
                    r1 = p.add_run(f"{k}: ")
                    r1.bold = True
                    r1.font.size = Pt(base_size)
                    r2 = p.add_run(str(v))
                    r2.font.size = Pt(base_size)
                    p.paragraph_format.space_after = Pt(1)

        elif el_type == "table":
            rows_data = _resolve(data, source) if source else []
            if isinstance(rows_data, list) and rows_data:
                columns = element.get("columns", [])
                _add_table(doc, rows_data, columns)

        elif el_type == "list":
            items = _resolve(data, source) if source else []
            fmt = element.get("item_format", "• {value}")
            if isinstance(items, list):
                for item in items:
                    if isinstance(item, dict):
                        txt = fmt.format(**item)
                    else:
                        txt = fmt.format(value=item)
                    _add_para(doc, txt, size=base_size - 1, after=2)

        elif el_type == "comments":
            items = _resolve(data, source) if source else []
            if isinstance(items, list):
                if items:
                    cp = doc.add_paragraph()
                    cr = cp.add_run(element.get("label", "Comments") + ":")
                    cr.bold = True
                    cr.font.size = Pt(base_size)
                    for c in items:
                        _add_para(doc, str(c), size=base_size - 1, after=4)
                else:
                    _add_para(doc, element.get("empty_text", "Nothing to report"), size=base_size - 1)


def _add_table(doc, rows: list[dict], columns: list[dict]):
    """Add a formatted table."""
    if not rows or not columns:
        return

    tbl = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    tbl.style = "Light Grid Accent 1"

    # Header
    for i, col in enumerate(columns):
        c = tbl.rows[0].cells[i]
        c.text = col.get("header", col.get("field", ""))
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
                rr.font.size = Pt(9)

    # Data
    for ri, row in enumerate(rows):
        for ci, col in enumerate(columns):
            val = row.get(col["field"], "")
            tbl.rows[ri + 1].cells[ci].text = str(val)

    # Font size
    for row in tbl.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(9)

    doc.add_paragraph()


def _add_styled_header(doc, text: str, color_hex: str = "#2F5496"):
    """Add a styled section header with dark background."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    pPr = p._p.get_or_add_pPr()
    fill = color_hex.lstrip("#")
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>'))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_para(doc, text: str, bold: bool = False, size: int = 10, align=None, after: int = 4):
    """Add a simple paragraph."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    r.font.size = Pt(size)
    r.bold = bold
    p.paragraph_format.space_after = Pt(after)
    return p


def _alignment(s: str):
    """Parse alignment string."""
    mapping = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT,
               "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    return mapping.get(s.lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _color(hex_str: str) -> RGBColor:
    """Parse hex color to RGBColor."""
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _resolve(data: dict, path: str) -> Any:
    """Resolve dot-path in nested dict."""
    parts = path.split(".")
    current = data
    for part in parts:
        if isinstance(current, dict):
            current = current.get(part)
        elif isinstance(current, list) and part.isdigit():
            idx = int(part)
            current = current[idx] if idx < len(current) else None
        else:
            return None
        if current is None:
            return None
    return current
